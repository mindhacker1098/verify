from flask import Flask, request, jsonify, send_file, abort
from flask_cors import CORS
from docx import Document
from docx2pdf import convert
import os
import logging
from pymongo import MongoClient
from gridfs import GridFS

app = Flask(__name__)
CORS(app)
app.config['PROPAGATE_EXCEPTIONS'] = True
app.config['JSONIFY_PRETTYPRINT_REGULAR'] = False

logging.basicConfig(level=logging.INFO)

# MongoDB setup
MONGO_URI = 'mongodb+srv://mindhacker1098:spn1098@cluster0.t66x9u5.mongodb.net/?retryWrites=true&w=majority'  # Replace with your MongoDB URI
client = MongoClient(MONGO_URI)
db = client['certificates_db']  # Replace with your database name
collection = db['certificates']  # Replace with your collection name
fs = GridFS(db)

enddates = {
    "1 month": "20-08-2024",
    "2 months": "20-09-2024",
    "3 months": "20-10-2024"
}

def replace_placeholder(doc, placeholder, value):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

def serialize_document(doc):
    """ Convert MongoDB document to a JSON serializable format """
    if doc:
        doc['_id'] = str(doc['_id'])  # Convert ObjectId to string
    return doc

@app.route('/verify/zidio/<string:id>', methods=['GET'])
def verify_certificate(id):
    certificate_id = f'zidio/{id}'
    logging.info(f"Verifying certificate_id: {certificate_id}")

    try:
        # Query MongoDB directly with the constructed ID
        certificate_data = collection.find_one({'certificate_id': certificate_id})

        # Convert document to a JSON serializable format
        certificate_data = serialize_document(certificate_data)

        if certificate_data:
            return jsonify({
                'valid': True,
                'certificate': certificate_data
            })
        else:
            return jsonify({'valid': False, 'message': 'Certificate not found'}), 404
    except Exception as e:
        logging.error(f"An error occurred: {str(e)}")
        return jsonify({'valid': False, 'error': str(e)}), 500

import pypandoc

@app.route('/download/zidio/<int:id>', methods=['GET'])
def download_certificate(id):
    certificate_id = f"zidio/{id:05d}"  # Format ID with leading zeros
    print(certificate_id)

    try:
        # Fetch metadata from MongoDB
        certificate_data = collection.find_one({'certificate_id': certificate_id})
        if not certificate_data:
            abort(404, description="Certificate not found")

        # Extract fields from MongoDB document
        name = certificate_data.get("name", "Unknown")
        position = certificate_data.get("position", "Unknown")
        start_date = certificate_data.get("start_date", "N/A")
        end_date = certificate_data.get("end_date", "N/A")

        input_docx_path = "offer.docx"  # Path to your DOCX template
        doc = Document(input_docx_path)

        replace_placeholder(doc, "<<Name>>", name)
        replace_placeholder(doc, "Position", position)
        replace_placeholder(doc, "<<start>>", start_date)
        replace_placeholder(doc, "<<end>>", end_date)

        temp_docx_path = f"temp.docx"
        doc.save(temp_docx_path)

        output_pdf_path = f"certificates/{name}.pdf"

        # Convert DOCX to PDF using pypandoc
        pypandoc.convert_file(temp_docx_path, 'pdf', outputfile=output_pdf_path)

        os.remove(temp_docx_path)

        if os.path.exists(output_pdf_path):
            return send_file(output_pdf_path, as_attachment=True, download_name='Certificate.pdf')
        else:
            logging.error(f"Generated file not found: {output_pdf_path}")
            abort(404)

    except Exception as e:
        logging.error(f"An error occurred: {str(e)}")
        return jsonify({'error': str(e)}), 500
if __name__ == '__main__':
    from waitress import serve
    serve(app, host='0.0.0.0', port=5000, connection_limit=10000, expose_tracebacks=True, ident=None, threads=4, url_scheme='http', asyncore_use_poll=True, cleanup_interval=30)
