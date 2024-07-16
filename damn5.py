from docx import Document
from docx2pdf import convert
import os
import pandas as pd
enddates={"1 month":"20-08-2024","2 months":"20-09-2024","3 months":"20-10-2024"}
def replace_placeholder_in_textboxes(doc, placeholder, value):
    for shape in doc.inline_shapes:
        if shape.text:
            if placeholder in shape.text:
                shape.text = shape.text.replace(placeholder, value)

    for textbox in doc.inline_shapes:
        if textbox.text_frame:
            for paragraph in textbox.text_frame.paragraphs:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

def replace_placeholder(doc, placeholder, value):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run.text)
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

    replace_placeholder_in_textboxes(doc, placeholder, value)
    

if __name__ == "__main__":
    df = pd.read_excel("part1/india_data.xlsx")
    i=0

  
    for index, row in df.iterrows():
        # if i<=2554:
        #     i+=1
        #     continue
        
        if index<=1999:
            continue

       
        name = row["Full Name (Will printed on certificate)"]
        position = row["Preferred Internship Program"]

        if position=="Mobile App Development(Flutter)":
            position="Mobile App Development"

        if position=="Data Science & Analytcs":
            position="Data Science & Analytics"

        duration = row["Internship Duration"]
        
        print(f"Name: {name}, Position: {position}")
        input_docx_path = "offer.docx"

  
        doc = Document(input_docx_path)

       
        replace_placeholder(doc,"<<Name>>", name)
        replace_placeholder(doc, "Position", position)
        replace_placeholder(doc, "<<start>>", "20-07-2024")
        replace_placeholder(doc, "<<end>>", enddates[duration])
        # replace_placeholder(doc, "<<issue>>", "12-06-2024")
        # replace_placeholder(doc, "<<Id>>", f"{index + 1:05d}")
        # replace_placeholder(doc, "<<Duration>>", duration)

        temp_docx_path = f"{index+2}"+".docx"
        doc.save(temp_docx_path)

      
        output_pdf_path = "certificate5/"+f"{index+2}"+".pdf"
        convert(temp_docx_path, output_pdf_path)

        
        os.remove(temp_docx_path)

        print(f"Conversion completed. PDF saved to: {output_pdf_path}")








#InternHiring #TechInternship #ProgrammingIntern #SoftwareInternship #DevelopmentInternship #InternshipOpportunity #TechJobs #CareerOpportunity #JobListing #EntryLevelDev #HiringNow #InternsWanted #CodeInternship #SoftwareDevelopment #JobHunt #JobAlert #JobSearch #JobOpportunity #JobPosting #JobHiring #EmploymentOpportunity #JobHunter #NewJob #GetHired #JobLink #TechOpportunity #NowHiring #CareerChange #JobMarket #ApplyNow #JobOffer #JobApplication #JoinOurTeam #CodeLife #TechCareers #SoftwareJobs #ProgrammingLife #DevelopersWanted #PromoteYourInternship #TechIntern #ComputerScienceInternship #CSInternship #SoftwareEngineerIntern #DataScienceInternship #AIInternship #MachineLearningInternship #CyberSecurityInternship #WebDevelopmentInternship #FrontendInternship #BackendInternship #FullstackInternship #PythonInternship #JavaInternship #NodeJSInternship #DevOpsInternship #MobileAppInternship #UIUXInternship #CloudComputingInternship #BlockchainInternship #BigDataInternship #ARVRInternship #GameDevInternship #DatabaseInternship #NetworkEngineeringInternship #RoboticsInternship #SoftwareTestingInternship #ITInternship #EmbeddedSystemsInternship #DataAnalysisInternship #BusinessIntelligenceInternship #ITSupportInternship #SysAdminInternship #TechSupportInternship #DevInternship #FrontendSkills #BackendSkills #DataScienceSkills #MachineLearningSkills #CyberSecuritySkills #AIEngineerIntern
     
    