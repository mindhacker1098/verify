from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx2pdf import convert

import os
import fitz
app = Flask(__name__)
CORS(app)
valid_certificates = {
    'CERT123': {'name': 'John Doe', 'pdf': 'certificates/CERT123.pdf'},
    'CERT456': {'name': 'Jane Smith', 'pdf': 'certificates/CERT456.pdf'},
    'CERT789': {'name': 'Alice Johnson', 'pdf': 'certificates/CERT789.pdf'}
}
enddates={"1 month":"20-08-2024","2 months":"20-09-2024","3 months":"20-10-2024"}
def replace_placeholder_in_textboxes(doc, placeholder, value):
    for shape in doc.inline_shapes:
        if shape.text:
            if placeholder in shape.text:
                shape.text = shape.text.replace(placeholder, value)

    for textbox in doc.inline_shapes:
        if textbox.text_frame:
            for paragraph in textbox.text_frame.paragraphs:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

def replace_placeholder(doc, placeholder, value):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run.text)
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

    replace_placeholder_in_textboxes(doc, placeholder, value)
def convert_pdf_to_image(pdf_path, output_folder='static/images'):
    pdf_document = fitz.open(pdf_path)
    image_paths = []


    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()
        image_path = os.path.join(output_folder, f"page_{page_num + 1}.png")
        pix.save(image_path)
        image_paths.append(image_path)

    return image_paths
   
@app.route('/verify', methods=['POST'])
def verify_certificate():
    data = request.get_json()
    certificate_id = data.get('certificateId')

    if certificate_id in valid_certificates:
        certificate = valid_certificates[certificate_id]
        pdf_path = certificate['pdf']
        image_paths = "CERT123.pdf"
        # image_paths = convert_pdf_to_image(pdf_path)
        name = "shakti"
        position = "web development"

        if position=="Mobile App Development(Flutter)":
            position="Mobile App Development"

        if position=="Data Science & Analytcs":
            position="Data Science & Analytics"

        duration = "3 months"
        
        print(f"Name: {name}, Position: {position}")
        input_docx_path = "offer5.docx"

  
        doc = Document(input_docx_path)

       
        replace_placeholder(doc,"<<Name>>", name)
        replace_placeholder(doc, "Position", position)
        replace_placeholder(doc, "<<start>>", "20-07-2024")
        replace_placeholder(doc, "<<end>>", enddates[duration])
        # replace_placeholder(doc, "<<issue>>", "12-06-2024")
        # replace_placeholder(doc, "<<Id>>", f"{index + 1:05d}")
        # replace_placeholder(doc, "<<Duration>>", duration)

        temp_docx_path = f"temp"+".docx"
        doc.save(temp_docx_path)

      
        output_pdf_path = "certificates/"+f"{name}"+".pdf"
        convert(temp_docx_path, output_pdf_path)

        
        os.remove(temp_docx_path)

        print(f"Conversion completed. PDF saved to: {output_pdf_path}")


        # image_paths = "CERT123.pdf"
        image_urls = [f"/certificate_image/{os.path.basename(image_path)}" for image_path in image_paths]
        return jsonify({
            'valid': True,
            'name': certificate['name'],
            'images': image_urls
        })
    else:
        return jsonify({'valid': False}), 404
       


if __name__ == '__main__':
    app.run(debug=True)
